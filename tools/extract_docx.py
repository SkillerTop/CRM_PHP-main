from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def iter_blocks(parent: DocumentObject):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_data(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([
            "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            for cell in row.cells
        ])
    return rows


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("document", type=Path)
    parser.add_argument("--start", type=int, default=1)
    parser.add_argument("--end", type=int)
    parser.add_argument("--format", choices=("json", "text"), default="json")
    args = parser.parse_args()

    path = args.document
    document = Document(path)
    blocks: list[dict[str, object]] = []

    for index, block in enumerate(iter_blocks(document), start=1):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                blocks.append({
                    "index": index,
                    "type": "paragraph",
                    "style": block.style.name if block.style else None,
                    "text": text,
                })
        else:
            blocks.append({
                "index": index,
                "type": "table",
                "rows": table_data(block),
            })

    blocks = [
        block for block in blocks
        if args.start <= int(block["index"]) <= (args.end or 10**9)
    ]
    payload = {
        "source": str(path),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "blocks": blocks,
    }
    if args.format == "json":
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return

    print(
        f"SOURCE: {path}\n"
        f"PARAGRAPHS: {len(document.paragraphs)}; TABLES: {len(document.tables)}; "
        f"SECTIONS: {len(document.sections)}"
    )
    for block in blocks:
        if block["type"] == "paragraph":
            print(f'\n[{block["index"]}] {block["style"]}: {block["text"]}')
            continue
        print(f'\n[{block["index"]}] TABLE')
        for row in block["rows"]:
            print(" | ".join(str(value).replace("\n", " / ") for value in row))


if __name__ == "__main__":
    main()
